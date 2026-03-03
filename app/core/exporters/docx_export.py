"""Exporter DOCX — dokumen Word untuk laporan skripsi/sidang."""

from __future__ import annotations

import logging
from pathlib import Path

from app.core.engines.base import TranscriptResult

logger = logging.getLogger(__name__)


def export_docx(result: TranscriptResult, output_path: str | Path) -> Path:
    """Export transkrip ke file .docx (Microsoft Word).

    Args:
        result: Hasil transkrip.
        output_path: Path file output.

    Returns:
        Path ke file yang dibuat.

    Raises:
        ImportError: Jika python-docx belum terinstall.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise ImportError("python-docx belum terinstall. Jalankan: pip install python-docx") from e

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Transkrip Wawancara", level=1)

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Bahasa: {result.language}\n").bold = False
    meta.add_run(f"Durasi: {_format_duration(result.duration)}\n")
    meta.add_run(f"Engine: {result.engine_name}\n")
    meta.add_run(f"Model: {result.model_name}\n")

    doc.add_paragraph("")

    current_speaker = ""
    for seg in result.segments:
        if seg.speaker and seg.speaker != current_speaker:
            doc.add_heading(seg.speaker, level=2)
            current_speaker = seg.speaker

        p = doc.add_paragraph()
        ts_run = p.add_run(f"[{_sec(seg.start)} → {_sec(seg.end)}] ")
        ts_run.bold = True
        ts_run.font.size = Pt(9)
        p.add_run(seg.text)

    doc.save(str(output_path))
    logger.info("DOCX disimpan ke %s", output_path)
    return output_path


def _sec(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    return f"{m:02d}:{s:02d}"


def _format_duration(seconds: float) -> str:
    h, remainder = divmod(int(seconds), 3600)
    m, s = divmod(remainder, 60)
    if h > 0:
        return f"{h}j {m}m {s}d"
    return f"{m}m {s}d"
